import os
import glob
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown
from bs4 import BeautifulSoup
import re
import requests
from urllib.parse import urlparse, urljoin
import mimetypes

class MarkdownToWordConverter:
    def __init__(self, base_dir=None):
        self.doc = Document()
        self.base_dir = Path(base_dir) if base_dir else Path.cwd()
        self.setup_styles()

    def setup_styles(self):
        """Set up custom styles for the document"""
        styles = self.doc.styles

        # Create heading styles if they don't exist
        for i in range(1, 7):
            try:
                heading_style = styles[f'Heading {i}']
            except KeyError:
                heading_style = styles.add_style(f'Heading {i}', 1)
                heading_style.base_style = styles['Normal']

    def download_image(self, url):
        """Download image from URL and return the content"""
        try:
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            return response.content
        except Exception as e:
            print(f"Failed to download image {url}: {e}")
            return None

    def get_image_path(self, src, current_file_dir=None):
        """Resolve image path - handles relative paths, absolute paths, and URLs"""
        if src.startswith(('http://', 'https://')):
            return src, 'url'

        # Handle relative paths
        if current_file_dir:
            image_path = current_file_dir / src
        else:
            image_path = self.base_dir / src

        if image_path.exists():
            return str(image_path), 'file'

        # Try relative to base directory if not found
        alt_path = self.base_dir / src
        if alt_path.exists():
            return str(alt_path), 'file'

        return None, None

    def add_image_to_doc(self, src, alt_text="", current_file_dir=None):
        """Add image to Word document"""
        image_path, path_type = self.get_image_path(src, current_file_dir)

        if not image_path:
            # Add placeholder text if image not found
            p = self.doc.add_paragraph(f"[Image not found: {src}]")
            run = p.runs[0]
            run.italic = True
            return

        try:
            if path_type == 'url':
                # Download image from URL
                image_data = self.download_image(image_path)
                if not image_data:
                    p = self.doc.add_paragraph(f"[Could not download image: {src}]")
                    run = p.runs[0]
                    run.italic = True
                    return

                # Save temporarily to add to document
                temp_path = self.base_dir / f"temp_image_{hash(src)}.jpg"
                with open(temp_path, 'wb') as f:
                    f.write(image_data)

                # Add image to document
                paragraph = self.doc.add_paragraph()
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

                # Determine appropriate size (max width 6 inches)
                run.add_picture(str(temp_path), width=Inches(6))

                # Clean up temp file
                temp_path.unlink()

            else:
                # Local file
                paragraph = self.doc.add_paragraph()
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()

                # Determine appropriate size (max width 6 inches)
                run.add_picture(image_path, width=Inches(6))

            # Add alt text as caption if provided
            if alt_text:
                caption_para = self.doc.add_paragraph(alt_text)
                caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in caption_para.runs:
                    run.italic = True

        except Exception as e:
            print(f"Error adding image {src}: {e}")
            # Add error placeholder
            p = self.doc.add_paragraph(f"[Error loading image: {src}]")
            run = p.runs[0]
            run.italic = True
        """Parse markdown content and convert to structured data"""
        # Convert markdown to HTML
        html = markdown.markdown(md_content, extensions=['extra', 'codehilite'])
        soup = BeautifulSoup(html, 'html.parser')

        return soup

    def add_element_to_doc(self, element, current_file_dir=None):
        """Add HTML element to Word document"""
        if element.name == 'h1':
            p = self.doc.add_heading(element.get_text().strip(), level=1)
        elif element.name == 'h2':
            p = self.doc.add_heading(element.get_text().strip(), level=2)
        elif element.name == 'h3':
            p = self.doc.add_heading(element.get_text().strip(), level=3)
        elif element.name == 'h4':
            p = self.doc.add_heading(element.get_text().strip(), level=4)
        elif element.name == 'h5':
            p = self.doc.add_heading(element.get_text().strip(), level=5)
        elif element.name == 'h6':
            p = self.doc.add_heading(element.get_text().strip(), level=6)
        elif element.name == 'p':
            # Check for images in paragraph
            images = element.find_all('img')
            if images:
                # Handle paragraph with images
                for img in images:
                    src = img.get('src', '')
                    alt = img.get('alt', '')
                    if src:
                        self.add_image_to_doc(src, alt, current_file_dir)

                # Add text content if any (excluding image tags)
                text_content = element.get_text().strip()
                if text_content:
                    p = self.doc.add_paragraph()
                    self.add_formatted_text(p, element, current_file_dir)
            elif element.get_text().strip():
                p = self.doc.add_paragraph()
                self.add_formatted_text(p, element, current_file_dir)
        elif element.name == 'img':
            # Standalone image
            src = element.get('src', '')
            alt = element.get('alt', '')
            if src:
                self.add_image_to_doc(src, alt, current_file_dir)
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                # Check for images in list items
                images = li.find_all('img')
                for img in images:
                    src = img.get('src', '')
                    alt = img.get('alt', '')
                    if src:
                        self.add_image_to_doc(src, alt, current_file_dir)

                text_content = li.get_text().strip()
                if text_content:
                    p = self.doc.add_paragraph(text_content, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                # Check for images in list items
                images = li.find_all('img')
                for img in images:
                    src = img.get('src', '')
                    alt = img.get('alt', '')
                    if src:
                        self.add_image_to_doc(src, alt, current_file_dir)

                text_content = li.get_text().strip()
                if text_content:
                    p = self.doc.add_paragraph(text_content, style='List Number')
        elif element.name == 'blockquote':
            p = self.doc.add_paragraph(element.get_text().strip())
            p.style = 'Quote'
        elif element.name == 'pre' or element.name == 'code':
            # Handle code blocks
            code_text = element.get_text()
            p = self.doc.add_paragraph(code_text)
            # Apply monospace font to code
            for run in p.runs:
                run.font.name = 'Courier New'
        elif element.name == 'hr':
            # Add a horizontal line (paragraph with bottom border)
            p = self.doc.add_paragraph()
            p.paragraph_format.border_bottom.color.rgb = None

    def add_formatted_text(self, paragraph, element, current_file_dir=None):
        """Add formatted text to paragraph, handling bold, italic, etc."""
        for content in element.contents:
            if hasattr(content, 'name'):
                if content.name == 'strong' or content.name == 'b':
                    run = paragraph.add_run(content.get_text())
                    run.bold = True
                elif content.name == 'em' or content.name == 'i':
                    run = paragraph.add_run(content.get_text())
                    run.italic = True
                elif content.name == 'code':
                    run = paragraph.add_run(content.get_text())
                    run.font.name = 'Courier New'
                elif content.name == 'a':
                    run = paragraph.add_run(content.get_text())
                    run.font.color.rgb = None  # Default hyperlink color
                elif content.name == 'img':
                    # Handle inline images
                    src = content.get('src', '')
                    alt = content.get('alt', '')
                    if src:
                        # For inline images, we'll add them after the paragraph
                        # Note: This is a limitation of python-docx - true inline images are complex
                        pass
                else:
                    paragraph.add_run(content.get_text())
            else:
                # Plain text
                paragraph.add_run(str(content))

    def process_markdown_file(self, file_path):
        """Process a single markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()

            # Add file name as a heading
            file_name = Path(file_path).stem
            self.doc.add_heading(f"Content from: {file_name}", level=1)

            # Parse and add content
            soup = self.parse_markdown_content(md_content)
            current_file_dir = Path(file_path).parent

            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'blockquote', 'pre', 'code', 'hr', 'img']):
                self.add_element_to_doc(element, current_file_dir)

            # Add page break after each file (except the last one)
            self.doc.add_page_break()

            print(f"Processed: {file_path}")

        except Exception as e:
            print(f"Error processing {file_path}: {str(e)}")

    def parse_index_file(self, index_path):
        """Parse index.md file to get chapter structure"""
        try:
            with open(index_path, 'r', encoding='utf-8') as f:
                index_content = f.read()

            # Convert markdown to HTML for easier parsing
            html = markdown.markdown(index_content, extensions=['extra'])
            soup = BeautifulSoup(html, 'html.parser')

            chapters = []
            current_chapter = None

            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li', 'a']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # New chapter/section
                    if current_chapter:
                        chapters.append(current_chapter)

                    level = int(element.name[1])
                    current_chapter = {
                        'title': element.get_text().strip(),
                        'level': level,
                        'files': [],
                        'subsections': []
                    }

                elif element.name == 'a' and element.get('href'):
                    # Found a link, likely to a markdown file
                    href = element.get('href')
                    title = element.get_text().strip()

                    if href.endswith('.md'):
                        if current_chapter:
                            current_chapter['files'].append({
                                'file': href,
                                'title': title
                            })
                        else:
                            # File without a chapter header
                            chapters.append({
                                'title': title,
                                'level': 1,
                                'files': [{'file': href, 'title': title}],
                                'subsections': []
                            })

            # Add the last chapter
            if current_chapter:
                chapters.append(current_chapter)

            return chapters

        except Exception as e:
            print(f"Error parsing index file: {str(e)}")
            return []

    def process_structured_content(self, base_dir, chapters):
        """Process markdown files according to the structure defined in index.md"""
        base_path = Path(base_dir)

        for chapter in chapters:
            # Add chapter heading
            if chapter['title']:
                self.doc.add_heading(chapter['title'], level=chapter['level'])

            # Process files in this chapter
            for file_info in chapter['files']:
                file_path = base_path / file_info['file']

                if not file_path.exists():
                    print(f"Warning: File not found: {file_path}")
                    continue

                # Add file title as subheading if different from chapter title
                if file_info['title'] != chapter['title'] and len(chapter['files']) > 1:
                    self.doc.add_heading(file_info['title'], level=min(chapter['level'] + 1, 6))

                # Process the markdown file content
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        md_content = f.read()

                    # Parse and add content
                    soup = self.parse_markdown_content(md_content)
                    current_file_dir = Path(file_path).parent

                    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'blockquote', 'pre', 'code', 'hr', 'img']):
                        # Adjust heading levels to fit within the chapter structure
                        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                            original_level = int(element.name[1])
                            new_level = min(chapter['level'] + original_level, 6)
                            self.doc.add_heading(element.get_text().strip(), level=new_level)
                        else:
                            self.add_element_to_doc(element, current_file_dir)

                    print(f"Processed: {file_path}")

                except Exception as e:
                    print(f"Error processing {file_path}: {str(e)}")

            # Add some spacing between chapters
            if chapter != chapters[-1]:  # Not the last chapter
                self.doc.add_paragraph()  # Add spacing
        """Convert all markdown files in a directory to a Word document"""
        input_path = Path(input_dir)

        if not input_path.exists():
            raise FileNotFoundError(f"Directory not found: {input_dir}")

        # Find all markdown files
        md_files = list(input_path.glob(pattern))

        if not md_files:
            raise FileNotFoundError(f"No markdown files found matching pattern '{pattern}' in {input_dir}")

        # Sort files alphabetically
        md_files.sort()

        print(f"Found {len(md_files)} markdown files")

        # Add title page
        title_para = self.doc.add_heading('Markdown Collection', 0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.doc.add_paragraph(f"Generated from {len(md_files)} markdown files")
        self.doc.add_page_break()

        # Process each file
        for i, md_file in enumerate(md_files):
            self.process_markdown_file(md_file)

            # Remove the last page break for the final file
            if i == len(md_files) - 1:
                # Remove last paragraph if it's a page break
                if self.doc.paragraphs and self.doc.paragraphs[-1].text == '':
                    p = self.doc.paragraphs[-1]._element
                    p.getparent().remove(p)

        # Save the document
        self.doc.save(output_file)
        print(f"Word document saved as: {output_file}")

    def convert_files(self, file_list, output_file):
        """Convert specific markdown files to a Word document"""
        if not file_list:
            raise ValueError("No files provided")

        # Verify all files exist
        for file_path in file_list:
            if not Path(file_path).exists():
                raise FileNotFoundError(f"File not found: {file_path}")

        print(f"Converting {len(file_list)} markdown files")

        # Add title page
        title_para = self.doc.add_heading('Markdown Collection', 0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.doc.add_paragraph(f"Generated from {len(file_list)} markdown files")
        self.doc.add_page_break()

        # Process each file
        for i, md_file in enumerate(file_list):
            self.process_markdown_file(md_file)

            # Remove the last page break for the final file
            if i == len(file_list) - 1:
                if self.doc.paragraphs and self.doc.paragraphs[-1].text == '':
                    p = self.doc.paragraphs[-1]._element
                    p.getparent().remove(p)

        # Save the document
        self.doc.save(output_file)
        print(f"Word document saved as: {output_file}")

def parse_arguments():
    """Parse command line arguments"""
    parser = argparse.ArgumentParser(
        description='Convert markdown files to Microsoft Word document',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
    python markdown_to_word.py --dir ./docs --out manual.docx
    python markdown_to_word.py --dir /path/to/markdown --out book.docx

The directory must contain an index.md file that defines the structure.
Example index.md format:

# Book Title

## Chapter 1: Introduction
- [Getting Started](chapter1.md)
- [Basic Concepts](basics.md)

## Chapter 2: Advanced Topics
- [Advanced Features](advanced.md)
- [Best Practices](practices.md)

Links should point to .md files relative to the directory.
        '''
    )

    parser.add_argument(
        '--dir',
        required=True,
        help='Path to directory containing markdown files and index.md'
    )

    parser.add_argument(
        '--out',
        required=True,
        help='Output Word document filename (e.g., document.docx)'
    )

    return parser.parse_args()

def main():
    """Main function with command line interface"""
    args = parse_arguments()

    try:
        converter = MarkdownToWordConverter()
        converter.convert_directory_with_index(args.dir, args.out)
        print("Conversion completed successfully!")

    except FileNotFoundError as e:
        print(f"Error: {e}")
        return 1
    except Exception as e:
        print(f"Error during conversion: {e}")
        return 1

    return 0

if __name__ == "__main__":
    main()
