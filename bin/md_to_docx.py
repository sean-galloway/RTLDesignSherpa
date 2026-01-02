#!/usr/bin/env python3
"""
md_to_docx.py — Convert Markdown (single file or expanded index) to DOCX/PDF via Pandoc.

Features:
- --expand-index: parse an index.md and inline linked chapter .md files in order
- --skip-index-content: exclude index file content, only include linked chapters
- --title-page [path]: prepend a title page (auto if no path given)
- --assets-dir (repeatable): added to Pandoc --resource-path
- --pagebreak: insert page breaks between concatenated files
- --section-breaks: start each top-level section on a new page (PDF only)
- --toc: generate table of contents with clickable links (depth=3)
- --number-sections: add section numbering (1, 1.1, 1.1.1)
- PDF: hyperref enabled for clickable TOC, internal links, and PDF bookmarks
- Emoji mapping for PDF robustness
- Force a Unicode-friendly PDF engine (xelatex) by default
- Wavedrom .json "images": render to SVG (python-wavedrom or wavedrom-cli) or degrade to links
- Wavedrom inline blocks: ```wavedrom JSON blocks rendered to SVG
- Mermaid diagrams: ```mermaid blocks rendered to PNG via mmdc CLI
- Font controls for XeLaTeX/LuaLaTeX: --mainfont/--monofont/--sansfont/--mathfont

Required tools for diagram rendering:
- Mermaid: npm install -g @mermaid-js/mermaid-cli (provides 'mmdc')
- Wavedrom: pip install wavedrom OR npm install -g wavedrom-cli

Note on DOCX TOC: Word may require updating the TOC (Ctrl+A, F9) to populate entries.
"""

import argparse
import pathlib
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import date

# ---------------------------
# Config / Helpers
# ---------------------------

MD_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+\.md)(#[^)]+)?\)", re.IGNORECASE)
IMG_JSON_RE = re.compile(r'!\[([^\]]*)\]\(([^)]+\.json)\)')
IMG_RE = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
# Fenced code block pattern for mermaid diagrams
# Captures: (1) optional figure title line, (2) mermaid code content
MERMAID_BLOCK_RE = re.compile(
    r'(#{2,4}\s+(?:Figure\s+\d+[:\s]+)?[^\n]+\n+)?```mermaid\s*\n(.*?)\n```',
    re.DOTALL | re.IGNORECASE
)

# Fenced code block pattern for wavedrom diagrams (inline JSON)
# Captures: (1) optional figure title line, (2) wavedrom JSON content
WAVEDROM_BLOCK_RE = re.compile(
    r'(#{2,4}\s+(?:Figure\s+\d+[:\s]+)?[^\n]+\n+)?```wavedrom\s*\n(.*?)\n```',
    re.DOTALL | re.IGNORECASE
)

EMOJI_MAP = {
    "✅": "✓",
    "❌": "✗",
    "➕": "+",
    "➖": "−",
    "⚠️": "⚠",
    "ℹ️": "ℹ",
    "🛠️": "🛠",
    "📌": "",
    "🚧": "",
}

def log(*a, **k):
    if not k.pop("_quiet", False):
        print(*a, **k, file=sys.stderr)

def read_text(p: pathlib.Path) -> str:
    return p.read_text(encoding="utf-8")

def write_text(p: pathlib.Path, s: str):
    p.write_text(s, encoding="utf-8")

def strip_or_map_emoji(s: str) -> str:
    for k, v in EMOJI_MAP.items():
        s = s.replace(k, v)
    return s

# ---- Mermaid handling ----

def render_mermaid_block(code: str, tmp_img_dir: pathlib.Path, idx: int, title: str = None, quiet: bool = False) -> str:
    """
    Render a Mermaid diagram code block to PNG.
    Returns markdown image reference or fallback code block.

    Note: PNG is used instead of SVG because SVG output from headless Chrome
    often has font/text rendering issues when fonts aren't installed in the
    headless environment.
    """
    out_png = tmp_img_dir / f"mermaid_{idx}.png"
    alt_text = title if title else f"Diagram {idx}"

    # Try mmdc CLI (mermaid-cli from npm)
    mmdc = shutil.which("mmdc")
    if mmdc:
        tmp_mmd = tmp_img_dir / f"mermaid_{idx}.mmd"
        write_text(tmp_mmd, code)

        # Create puppeteer config to disable sandbox (required on Ubuntu 23.10+)
        puppeteer_config = tmp_img_dir / "puppeteer-config.json"
        if not puppeteer_config.exists():
            write_text(puppeteer_config, '{"args": ["--no-sandbox", "--disable-setuid-sandbox"]}')

        try:
            # Use PNG output with high scale for quality, white background for better PDF embedding
            cmd = [mmdc, "-i", str(tmp_mmd), "-o", str(out_png), "-b", "white",
                   "-p", str(puppeteer_config), "-s", "2"]  # scale=2 for higher resolution
            if quiet:
                cmd.append("-q")
            subprocess.run(cmd, check=True, capture_output=True, timeout=30)
            if out_png.exists():
                if not quiet:
                    log(f"  Rendered mermaid_{idx}.png: {alt_text}")
                return f"![{alt_text}]({out_png.as_posix()})"
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
            if not quiet:
                log(f"  Warning: Mermaid render failed for block {idx}: {e}")

    # Fallback: keep as code block
    if not quiet:
        log(f"  Warning: mmdc not found, keeping mermaid block {idx} as code")
    return f"```\n{code}\n```"

def rewrite_mermaid_blocks(md_text: str, tmp_img_dir: pathlib.Path, quiet: bool = False) -> str:
    """Replace all ```mermaid blocks with rendered PNG images."""
    idx = [0]  # Use list to allow mutation in nested function

    def _sub(m):
        title_line = m.group(1)  # Optional figure title line (e.g., "### Figure 1: Title\n")
        code = m.group(2).strip()
        idx[0] += 1

        # Extract clean title from the heading line
        title = None
        if title_line:
            # Remove markdown heading markers and clean up
            title = re.sub(r'^#+\s*', '', title_line.strip())
            # Remove "Figure N:" prefix since Pandoc adds its own
            title = re.sub(r'^Figure\s+\d+[:\s]*', '', title)
            title = title.rstrip(':').strip()

        return render_mermaid_block(code, tmp_img_dir, idx[0], title, quiet)

    return MERMAID_BLOCK_RE.sub(_sub, md_text)


# ---- Inline Wavedrom handling (fenced code blocks) ----

def render_wavedrom_inline(json_code: str, tmp_img_dir: pathlib.Path, idx: int, title: str = None, quiet: bool = False, register_low: bool = False) -> str:
    """
    Render an inline WaveDrom JSON code block to SVG.
    Returns markdown image reference or fallback code block.

    Args:
        register_low: If True, use raw LaTeX (not figure) and add to List of Waveforms
    """
    import json as json_module
    out_svg = tmp_img_dir / f"wavedrom_{idx}.svg"
    alt_text = title if title else f"Waveform {idx}"

    # Write JSON to temp file
    tmp_json = tmp_img_dir / f"wavedrom_{idx}.json"
    write_text(tmp_json, json_code)

    # Helper to build result - add LoW entry when register_low is True
    def _build_result(svg_path: pathlib.Path) -> str:
        if register_low:
            # Convert SVG to PDF for better LaTeX compatibility
            pdf_path = svg_path.with_suffix('.pdf')
            try:
                subprocess.run(
                    ['inkscape', str(svg_path), '--export-filename=' + str(pdf_path)],
                    check=True, capture_output=True
                )
            except (subprocess.CalledProcessError, FileNotFoundError):
                # Fallback to SVG if inkscape fails
                pdf_path = svg_path

            escaped_title = alt_text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
            img_path = pdf_path.as_posix()
            return f"""
```{{=latex}}
\\begin{{center}}
\\includegraphics[width=\\linewidth,keepaspectratio]{{{img_path}}}

\\vspace{{0.5em}}
\\textbf{{{escaped_title}}}
\\addcontentsline{{low}}{{waveform}}{{{escaped_title}}}
\\end{{center}}
```
"""
        # Standard markdown image (will become figure in LaTeX)
        return f"![{alt_text}]({svg_path.as_posix()})"

    # Try python-wavedrom first
    svg_text = try_render_wavedrom_python(tmp_json)
    if svg_text:
        write_text(out_svg, svg_text)
        if not quiet:
            log(f"  Rendered wavedrom_{idx}.svg (python-wavedrom): {alt_text}")
        return _build_result(out_svg)

    # Try wavedrom-cli
    if try_render_wavedrom_cli(tmp_json, out_svg):
        if not quiet:
            log(f"  Rendered wavedrom_{idx}.svg (wavedrom-cli): {alt_text}")
        return _build_result(out_svg)

    # Fallback: keep as code block
    if not quiet:
        log(f"  Warning: wavedrom render failed for block {idx}, keeping as code")
    return f"```json\n{json_code}\n```"


def rewrite_wavedrom_blocks(md_text: str, tmp_img_dir: pathlib.Path, quiet: bool = False, register_low: bool = False) -> str:
    """Replace all ```wavedrom blocks with rendered SVG images.

    Args:
        register_low: If True, add entries to List of Waveforms for titled diagrams
    """
    idx = [0]  # Use list to allow mutation in nested function

    def _sub(m):
        title_line = m.group(1)  # Optional figure title line (e.g., "### Figure 1: Title\n")
        code = m.group(2).strip()
        idx[0] += 1

        # Extract clean title from the heading line
        title = None
        if title_line:
            # Remove markdown heading markers and clean up
            title = re.sub(r'^#+\s*', '', title_line.strip())
            # Remove "Figure N:" prefix since Pandoc adds its own
            title = re.sub(r'^Figure\s+\d+[:\s]*', '', title)
            title = title.rstrip(':').strip()

        return render_wavedrom_inline(code, tmp_img_dir, idx[0], title, quiet, register_low)

    return WAVEDROM_BLOCK_RE.sub(_sub, md_text)

def collect_from_index(index_path: pathlib.Path, skip_index: bool = False) -> list[pathlib.Path]:
    """
    Scan the index markdown for links to .md files and return them in order.
    Keeps only links that resolve relative to the index folder.

    Args:
        index_path: Path to the index markdown file
        skip_index: If True, don't include the index file itself (only linked chapters)
    """
    root = index_path.parent
    content = read_text(index_path)
    links = []
    seen = set()

    # Inline comment include directive support:  <!-- include: path/to/file.md -->
    for inc in re.findall(r'<!--\s*include:\s*([^\s>]+\.md)\s*-->', content, flags=re.IGNORECASE):
        p = (root / inc).resolve()
        if p.exists() and p.suffix.lower() == ".md" and p not in seen:
            links.append(p); seen.add(p)

    for m in MD_LINK_RE.finditer(content):
        rel = m.group(2).strip()
        if "://" in rel or rel.startswith("#"):
            continue
        p = (root / rel).resolve()
        if p.suffix.lower() == ".md" and p.exists() and p not in seen:
            links.append(p); seen.add(p)

    # Only include index file if not skipping it
    if not skip_index:
        idx = index_path.resolve()
        if idx not in seen:
            links = [idx] + links
        else:
            links = [idx] + [p for p in links if p != idx]

    return links

def rewrite_image_paths_for_file(text: str, source_file: pathlib.Path) -> str:
    """
    Rewrite relative image paths in markdown text to absolute paths.
    This is necessary when concatenating markdown files from different directories.
    """
    def _sub(m):
        alt, path = m.group(1), m.group(2)
        # Skip absolute paths, URLs, and already processed paths
        if path.startswith('/') or '://' in path:
            return m.group(0)
        # Resolve relative path based on source file location
        abs_path = (source_file.parent / path).resolve()
        if abs_path.exists():
            return f"![{alt}]({abs_path.as_posix()})"
        # If file doesn't exist, keep original reference
        return m.group(0)
    return IMG_RE.sub(_sub, text)

def concat_markdown(files: list[pathlib.Path], pagebreak: bool) -> str:
    parts = []
    for i, f in enumerate(files):
        text = read_text(f).rstrip() + "\n"
        # Rewrite relative image paths to absolute paths
        text = rewrite_image_paths_for_file(text, f)
        parts.append(text)
        if pagebreak and i < len(files) - 1:
            parts.append('\n::: {.pagebreak}\n:::\n')
    return "\n".join(parts)

# ---- Wavedrom handling ----

def try_render_wavedrom_python(json_path: pathlib.Path) -> str | None:
    try:
        import wavedrom
        # Read JSON content from file
        json_content = json_path.read_text()
        # render() returns an svgwrite.Drawing object, not a string
        drawing = wavedrom.render(json_content)
        # Convert Drawing to SVG string
        if hasattr(drawing, 'tostring'):
            return drawing.tostring()
        elif isinstance(drawing, str):
            return drawing
        return None
    except Exception:
        return None

def try_render_wavedrom_cli(json_path: pathlib.Path, out_svg: pathlib.Path) -> bool:
    # Try direct wavedrom-cli first, then npx as fallback
    cli_cmd = None
    if shutil.which("wavedrom-cli"):
        cli_cmd = ["wavedrom-cli"]
    elif shutil.which("npx"):
        cli_cmd = ["npx", "wavedrom-cli"]
    else:
        return False
    try:
        subprocess.run(
            cli_cmd + ["-i", str(json_path), "-s", str(out_svg)],
            check=True,
            capture_output=True
        )
        return out_svg.exists()
    except subprocess.CalledProcessError:
        return False

def rewrite_wavedrom_images(md_text: str, base_dir: pathlib.Path, tmp_img_dir: pathlib.Path) -> str:
    def _sub(m):
        alt, rel = m.group(1), m.group(2)
        jp = (base_dir / rel).resolve()
        if not jp.exists():
            return f"[{alt or 'diagram (wavedrom)'}]({rel})"
        svg_text = try_render_wavedrom_python(jp)
        if svg_text:
            out_svg = tmp_img_dir / (jp.stem + ".svg")
            write_text(out_svg, svg_text)
            return f"![{alt}]({out_svg.as_posix()})"
        out_svg = tmp_img_dir / (jp.stem + ".svg")
        if try_render_wavedrom_cli(jp, out_svg):
            return f"![{alt}]({out_svg.as_posix()})"
        return f"[{alt or 'diagram (wavedrom)'}]({rel})"
    return IMG_JSON_RE.sub(_sub, md_text)

# ---------------------------
# Pandoc runners
# ---------------------------

def build_resource_path(args, input_path: pathlib.Path) -> str:
    paths = [str(pathlib.Path(a).resolve()) for a in args.assets_dir]
    paths.append(str(input_path.resolve().parent))
    seen = set()
    dedup = []
    for p in paths:
        if p not in seen:
            seen.add(p)
            dedup.append(p)
    return ":".join(dedup)

def run_pandoc_docx(md_path: pathlib.Path, out_docx: pathlib.Path, args):
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("pandoc not found on PATH")
    cmd = [
        pandoc, str(md_path), "-o", str(out_docx),
        "--standalone",
        "--from", "markdown+yaml_metadata_block+pipe_tables+fenced_code_blocks",
    ]
    if args.reference_doc:
        cmd += ["--reference-doc", args.reference_doc]
    if args.toc:
        cmd += ["--toc", "--toc-depth=3"]
    if args.number_sections:
        cmd += ["--number-sections"]
    if args.quiet:
        cmd += ["--quiet"]
    cmd += ["--resource-path", build_resource_path(args, pathlib.Path(args.input))]
    cmd += ["-V", "graphics=true"]  # help Pandoc with figures/pagebreaks
    if not args.quiet:
        log(f"DOCX command: {' '.join(cmd)}")
    subprocess.run(cmd, check=True)

def run_pandoc_pdf(md_path: pathlib.Path, out_pdf: pathlib.Path, args, tmp_dir: pathlib.Path = None, title_page_tex: str = None):
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("pandoc not found on PATH")
    engine = args.pdf_engine or "xelatex"
    cmd = [
        pandoc, str(md_path), "-o", str(out_pdf),
        "--standalone",
        "--from", "markdown+yaml_metadata_block+pipe_tables+fenced_code_blocks",
        f"--pdf-engine={engine}",
        "--resource-path", build_resource_path(args, pathlib.Path(args.input)),
    ]

    # Title page - use include-before-body to place BEFORE TOC
    if title_page_tex and tmp_dir:
        title_tex_file = tmp_dir / "title_page.tex"
        write_text(title_tex_file, title_page_tex)
        cmd += ["--include-before-body", str(title_tex_file)]

    if args.toc:
        cmd += ["--toc", "--toc-depth=3"]
    if args.number_sections:
        cmd += ["--number-sections"]
    if args.quiet:
        cmd += ["--quiet"]

    # --- Font controls: sensible defaults for wide Unicode coverage ---
    default_main = "DejaVu Serif"
    default_mono = "DejaVu Sans Mono"
    mainfont = args.mainfont or default_main
    monofont = args.monofont or default_mono
    cmd += ["-V", f"mainfont={mainfont}", "-V", f"monofont={monofont}"]
    if args.sansfont:
        cmd += ["-V", f"sansfont={args.sansfont}"]
    if args.mathfont:
        cmd += ["-V", f"mathfont={args.mathfont}"]

    # --- Hyperref configuration for clickable links and PDF bookmarks ---
    cmd += ["-V", "colorlinks=true"]
    cmd += ["-V", "linkcolor=blue"]
    cmd += ["-V", "toccolor=blue"]
    cmd += ["-V", "urlcolor=blue"]
    cmd += ["-V", "bookmarks=true"]
    cmd += ["-V", "bookmarksnumbered=true"]

    # Margin control
    if args.narrow_margins:
        cmd += ["-V", "geometry:margin=0.75in"]

    # Section breaks - start each top-level section on a new page
    if args.section_breaks and tmp_dir:
        # Use titlesec package to add \clearpage before each section
        header_tex = tmp_dir / "section_breaks.tex"
        write_text(header_tex, r"""
\usepackage{titlesec}
\newcommand{\sectionbreak}{\clearpage}
""")
        cmd += ["--include-in-header", str(header_tex)]

    # List of Waveforms - custom list type using tocloft package
    # Also include svg package for direct SVG inclusion (requires inkscape)
    if args.low and tmp_dir:
        low_header_tex = tmp_dir / "list_of_waveforms.tex"
        write_text(low_header_tex, r"""
\usepackage{tocloft}
\newlistof{waveform}{low}{List of Waveforms}
\newcommand{\listofwaveforms}{\listofwaveform}
\usepackage{svg}
\svgsetup{inkscapelatex=false}
""")
        cmd += ["--include-in-header", str(low_header_tex)]
        # svg package needs --shell-escape to call inkscape
        cmd += ["--pdf-engine-opt=--shell-escape"]

    if not args.quiet:
        log(f"PDF command: {' '.join(cmd)}")
    subprocess.run(cmd, check=True)

# ---------------------------
# Main
# ---------------------------

def parse_args():
    p = argparse.ArgumentParser(
        description="Convert Markdown (single file or expanded index) to DOCX/PDF via Pandoc."
    )
    p.add_argument("input", help="Input Markdown file (index or chapter).")
    p.add_argument("output", help="Output .docx filename (PDF uses same basename).")
    p.add_argument("--reference-doc", help="Path to a DOCX reference (template).")
    p.add_argument("--pdf", action="store_true", help="Also emit a PDF with same basename.")
    p.add_argument("--toc", action="store_true", help="Include a table of contents.")
    p.add_argument("--number-sections", action="store_true", help="Number sections (1, 1.1, 1.1.1, etc.).")
    p.add_argument("--quiet", action="store_true", help="Reduce Pandoc chatter.")
    p.add_argument("--pagebreak", action="store_true", help="Insert page breaks between concatenated files.")
    p.add_argument("--assets-dir", action="append", default=[],
                help="Add an assets/resource directory (repeatable).")
    p.add_argument("--title-page", nargs="?", const="__AUTO__",
                help="Prepend a title page. If value omitted, auto-generate; otherwise treat as path to a .md file.")
    p.add_argument("--expand-index", action="store_true",
                help="Parse the input index and inline linked chapter .md files in order.")
    p.add_argument("--skip-index-content", action="store_true",
                help="When using --expand-index, don't include the index file content (only chapters).")
    p.add_argument("--debug-md", action="store_true",
                help="Save the merged markdown file for debugging (as output.build.md).")
    p.add_argument("--no-mermaid", action="store_true",
                help="Skip mermaid rendering (leave as code blocks for faster generation).")
    p.add_argument("--no-wavedrom", action="store_true",
                help="Skip wavedrom rendering (leave as code blocks for faster generation).")
    p.add_argument("--narrow-margins", action="store_true",
                help="Use narrow margins (0.75in) for more content per page.")
    p.add_argument("--section-breaks", action="store_true",
                help="Start each top-level section on a new page (PDF only).")
    p.add_argument("--lot", action="store_true",
                help="Include List of Tables after TOC (PDF only).")
    p.add_argument("--lof", action="store_true",
                help="Include List of Figures after TOC (PDF only).")
    p.add_argument("--low", action="store_true",
                help="Include List of Waveforms after TOC (PDF only).")
    p.add_argument("--pdf-engine", default=None,
                help="Override PDF engine (default: xelatex).")
    # Font controls (XeLaTeX/LuaLaTeX)
    p.add_argument("--mainfont", default=None, help="Main text font (e.g., 'DejaVu Serif', 'Noto Serif').")
    p.add_argument("--monofont", default=None, help="Monospace font (e.g., 'DejaVu Sans Mono', 'Noto Sans Mono').")
    p.add_argument("--sansfont", default=None, help="Sans-serif font (optional).")
    p.add_argument("--mathfont", default=None, help="Math font (optional).")
    p.add_argument("--style", default=None,
                help="YAML style config file for corporate DOCX styling (post-processes DOCX before PDF).")
    return p.parse_args()

# ---------------------------
# DOCX Corporate Styling
# ---------------------------

def add_title_page_to_docx(doc, config: dict, colors: dict) -> None:
    """Add a title page to the beginning of a DOCX document.

    Creates a left-aligned title page matching corporate style:
    - Company name with split coloring (e.g., "Qernel" black + "AI" red)
    - Title and subtitle
    - Date
    - Page break before TOC

    Args:
        doc: python-docx Document object
        config: title_page config dict from YAML
        colors: colors dict from YAML for resolving color references
    """
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def hex_to_rgb(hex_color: str) -> RGBColor:
        hex_color = hex_color.lstrip('#')
        return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

    def resolve_color(color_name: str) -> str:
        if color_name.startswith('#'):
            return color_name
        return colors.get(color_name, colors.get('black', '#000000'))

    # Get title page content from config
    company = config.get('company', 'QernelAI')
    title = config.get('title', 'Document Title')
    subtitle = config.get('subtitle', '')
    date_str = config.get('date', '')
    company_color = resolve_color(config.get('company_color', 'primary'))
    title_color = resolve_color(config.get('title_color', 'secondary'))

    # Find the very first element to insert before (skip any existing TOC heading)
    body = doc.element.body
    first_element = body[0] if len(body) > 0 else None

    # We'll build paragraphs and insert them all at once at the beginning
    # Create a list to hold paragraphs in order
    title_paragraphs = []

    # Spacer paragraph at top (creates vertical space ~1/3 down page)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(200)  # Large space at top
    title_paragraphs.append(spacer)

    # Company name - special handling for "QernelAI" split coloring
    # "Qernel" in black, "AI" in red (primary color) - 72pt
    company_para = doc.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if company == "QernelAI":
        # Split into "Qernel" (black) + "AI" (red)
        run1 = company_para.add_run("Qernel")
        run1.font.size = Pt(72)
        run1.font.bold = True
        run1.font.name = 'Calibri'
        run1.font.color.rgb = hex_to_rgb('#000000')  # Black

        run2 = company_para.add_run("AI")
        run2.font.size = Pt(72)
        run2.font.bold = True
        run2.font.name = 'Calibri'
        run2.font.color.rgb = hex_to_rgb(company_color)  # Red
    else:
        run = company_para.add_run(company)
        run.font.size = Pt(72)
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.color.rgb = hex_to_rgb(company_color)
    title_paragraphs.append(company_para)

    # Title line (combines title and subtitle if present) - 36pt
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_before = Pt(20)
    title_text = title
    if subtitle:
        title_text = f"{title} {subtitle}"
    run = title_para.add_run(title_text)
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.color.rgb = hex_to_rgb('#000000')  # Black
    title_paragraphs.append(title_para)

    # Date - 36pt
    if date_str:
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_para.paragraph_format.space_before = Pt(20)
        run = date_para.add_run(date_str)
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.color.rgb = hex_to_rgb('#000000')
        title_paragraphs.append(date_para)

    # Page break after title page
    break_para = doc.add_paragraph()
    run = break_para.add_run()
    run.add_break(WD_BREAK.PAGE)
    title_paragraphs.append(break_para)

    # Insert all title page paragraphs at the beginning of the document
    # Insert in reverse order so they end up in correct order
    for para in reversed(title_paragraphs):
        if first_element is not None:
            first_element.addprevious(para._p)
        first_element = para._p


def apply_docx_style(docx_path: pathlib.Path, style_config_path: pathlib.Path, quiet: bool = False) -> bool:
    """Apply corporate styling to a DOCX file using a YAML config.

    Returns True if styling was applied, False if dependencies missing.
    """
    try:
        import yaml
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        if not quiet:
            log(f"  Warning: DOCX styling skipped (missing dependency: {e})")
        return False

    def hex_to_rgb(hex_color: str) -> RGBColor:
        hex_color = hex_color.lstrip('#')
        return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

    def resolve_color(color_name: str, colors: dict) -> str:
        if color_name.startswith('#'):
            return color_name
        return colors.get(color_name, colors.get('black', '#000000'))

    def add_shading(element, color_hex: str):
        color_hex = color_hex.lstrip('#')
        pPr = element.get_or_add_pPr() if hasattr(element, 'get_or_add_pPr') else element
        for shd in pPr.findall(qn('w:shd')):
            pPr.remove(shd)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        pPr.append(shd)

    def add_bottom_border(paragraph, color_hex: str, width: int = 12):
        color_hex = color_hex.lstrip('#')
        pPr = paragraph._p.get_or_add_pPr()
        for pBdr in pPr.findall(qn('w:pBdr')):
            pPr.remove(pBdr)
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(width))
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def set_cell_shading(cell, color_hex: str):
        color_hex = color_hex.lstrip('#')
        tcPr = cell._tc.get_or_add_tcPr()
        for shd in tcPr.findall(qn('w:shd')):
            tcPr.remove(shd)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    def get_heading_level(para) -> int:
        style_name = para.style.name if para.style else ''
        if style_name.startswith('Heading'):
            try:
                return int(style_name.split()[-1])
            except (ValueError, IndexError):
                pass
        return 0

    # Load config
    with open(style_config_path) as f:
        config = yaml.safe_load(f)

    colors = config.get('colors', {})
    fonts = config.get('fonts', {})
    headings = config.get('headings', {})
    body_config = config.get('body', {})
    table_config = config.get('tables', {})

    # Open document
    doc = Document(str(docx_path))

    # Add title page if configured
    title_page_config = config.get('title_page', {})
    if title_page_config.get('enabled', False):
        add_title_page_to_docx(doc, title_page_config, colors)
        if not quiet:
            log(f"  Added title page to DOCX")

    # Process paragraphs
    for para in doc.paragraphs:
        level = get_heading_level(para)

        if level > 0:
            h_key = f'h{level}'
            if h_key in headings:
                h_cfg = headings[h_key]
                # Apply font styling to runs
                for run in para.runs:
                    if 'font_size' in h_cfg:
                        run.font.size = Pt(h_cfg['font_size'])
                    if 'color' in h_cfg:
                        color = resolve_color(h_cfg['color'], colors)
                        run.font.color.rgb = hex_to_rgb(color)
                    if h_cfg.get('bold'):
                        run.font.bold = True
                    if h_cfg.get('italic'):
                        run.font.italic = True

                # Paragraph formatting
                pf = para.paragraph_format
                if 'space_before' in h_cfg:
                    pf.space_before = Pt(h_cfg['space_before'])
                if 'space_after' in h_cfg:
                    pf.space_after = Pt(h_cfg['space_after'])

                # Background shading
                if 'background' in h_cfg:
                    bg_color = resolve_color(h_cfg['background'], colors)
                    add_shading(para._p, bg_color)

                # Underline border
                if 'underline_color' in h_cfg:
                    ul_color = resolve_color(h_cfg['underline_color'], colors)
                    add_bottom_border(para, ul_color)

        elif para.style and para.style.name in ['Normal', 'Body Text', 'First Paragraph']:
            for run in para.runs:
                if body_config.get('font_size'):
                    run.font.size = Pt(body_config['font_size'])

    # Process tables - style header row
    header_config = table_config.get('header', {})
    for table in doc.tables:
        if table.rows:
            for cell in table.rows[0].cells:
                if header_config.get('background'):
                    set_cell_shading(cell, header_config['background'])
                for para in cell.paragraphs:
                    for run in para.runs:
                        if header_config.get('bold'):
                            run.font.bold = True
                        if header_config.get('font_size'):
                            run.font.size = Pt(header_config['font_size'])

    # Add footer with confidential text
    hf_config = config.get('header_footer', {})
    if hf_config.get('enabled', False):
        company = config.get('company', {})
        confidential = company.get('confidential_text', '')
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            if footer.paragraphs:
                para = footer.paragraphs[0]
                para.clear()
            else:
                para = footer.add_paragraph()
            para.text = confidential
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(hf_config.get('font_size', 8))

    # Save
    doc.save(str(docx_path))
    if not quiet:
        log(f"  Applied corporate styling from {style_config_path.name}")
    return True

def update_docx_toc(docx_path: pathlib.Path, soffice_path: str, quiet: bool = False) -> bool:
    """Update Table of Contents in a DOCX file using LibreOffice UNO.

    Uses LibreOffice's UNO API to open the document, refresh all indexes,
    and save it. This populates the TOC with actual entries.

    Args:
        docx_path: Path to the DOCX file
        soffice_path: Path to soffice/libreoffice executable
        quiet: Suppress output messages

    Returns:
        True if TOC was updated successfully, False otherwise
    """
    abs_docx = str(docx_path.resolve())

    # Create a standalone script that uses system Python (which has uno)
    # This is needed because venv Python typically doesn't have access to uno
    uno_script = f'''#!/usr/bin/env python3
import sys
import time
import subprocess

def main():
    try:
        import uno
        from com.sun.star.beans import PropertyValue
    except ImportError:
        print("ERROR: uno module not available", file=sys.stderr)
        return 1

    soffice = "{soffice_path}"
    docx_file = "{abs_docx}"

    # Start LibreOffice in listening mode
    proc = subprocess.Popen([
        soffice,
        "--headless",
        "--invisible",
        "--nofirststartwizard",
        "--accept=socket,host=localhost,port=2002;urp;StarOffice.ServiceManager"
    ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    time.sleep(3)  # Wait for LibreOffice to start

    try:
        local_context = uno.getComponentContext()
        resolver = local_context.ServiceManager.createInstanceWithContext(
            "com.sun.star.bridge.UnoUrlResolver", local_context)

        # Try to connect (with retries)
        ctx = None
        for attempt in range(10):
            try:
                ctx = resolver.resolve(
                    "uno:socket,host=localhost,port=2002;urp;StarOffice.ComponentContext")
                break
            except:
                time.sleep(1)

        if ctx is None:
            print("ERROR: Could not connect to LibreOffice", file=sys.stderr)
            proc.terminate()
            return 1

        smgr = ctx.ServiceManager
        desktop = smgr.createInstanceWithContext("com.sun.star.frame.Desktop", ctx)

        # Open document hidden
        file_url = "file://" + docx_file
        load_props = (PropertyValue(Name="Hidden", Value=True),)
        doc = desktop.loadComponentFromURL(file_url, "_blank", 0, load_props)

        if doc:
            # Update all indexes
            indexes = doc.getDocumentIndexes()
            for i in range(indexes.getCount()):
                indexes.getByIndex(i).update()

            # Save
            doc.store()
            doc.close(True)
            print("TOC updated successfully")

        # Shutdown LibreOffice
        desktop.terminate()
        proc.wait(timeout=10)
        return 0

    except Exception as e:
        print(f"ERROR: {{e}}", file=sys.stderr)
        try:
            proc.terminate()
        except:
            pass
        return 1

if __name__ == "__main__":
    sys.exit(main())
'''

    # Write script to temp file and run with system Python
    with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False) as f:
        f.write(uno_script)
        script_path = f.name

    try:
        # Use system Python which has uno module (not venv Python)
        system_python = "/usr/bin/python3"
        if not pathlib.Path(system_python).exists():
            system_python = shutil.which("python3")

        result = subprocess.run(
            [system_python, script_path],
            capture_output=True,
            text=True,
            timeout=120
        )

        if result.returncode == 0:
            if not quiet:
                log("  TOC updated successfully")
            return True
        else:
            if not quiet:
                stderr = result.stderr.strip() if result.stderr else "unknown error"
                log(f"  Warning: TOC update failed: {stderr}")
            return False

    except subprocess.TimeoutExpired:
        if not quiet:
            log("  Warning: TOC update timed out")
        return False
    except Exception as e:
        if not quiet:
            log(f"  Warning: TOC update failed: {e}")
        return False
    finally:
        # Clean up temp script
        try:
            pathlib.Path(script_path).unlink()
        except:
            pass


def reapply_title_page_fonts(docx_path: pathlib.Path, style_config_path: pathlib.Path, quiet: bool = False) -> bool:
    """Re-apply title page font sizes after LibreOffice TOC update.

    LibreOffice may reset font sizes when updating the TOC. This function
    finds the title page paragraphs and reapplies the correct font sizes.

    Args:
        docx_path: Path to the DOCX file
        style_config_path: Path to the YAML style config
        quiet: Suppress output messages

    Returns:
        True if successful, False otherwise
    """
    try:
        import yaml
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        return False

    def hex_to_rgb(hex_color: str) -> RGBColor:
        hex_color = hex_color.lstrip('#')
        return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

    # Load config
    with open(style_config_path) as f:
        config = yaml.safe_load(f)

    title_page_config = config.get('title_page', {})
    if not title_page_config.get('enabled', False):
        return True

    colors = config.get('colors', {})
    company = title_page_config.get('company', 'QernelAI')
    company_color = colors.get(title_page_config.get('company_color', 'primary'), '#CC0000')

    doc = Document(str(docx_path))

    # Find title page paragraphs (before the first page break or TOC)
    # Title page has: spacer, company name, title, date, page break
    title_page_end = 0
    for i, para in enumerate(doc.paragraphs):
        # Check if this is a page break or TOC heading
        if 'Table of Contents' in para.text or para.style.name.startswith('TOC'):
            title_page_end = i
            break
        # Check for page break in runs
        for run in para.runs:
            if run._element.xml.find('w:br') != -1 and 'w:type="page"' in run._element.xml:
                title_page_end = i + 1
                break
        if title_page_end > 0:
            break

    if title_page_end == 0:
        title_page_end = min(5, len(doc.paragraphs))

    # Apply font sizes to title page paragraphs
    for i, para in enumerate(doc.paragraphs[:title_page_end]):
        text = para.text.strip()
        if not text:
            continue

        # Determine what this paragraph is based on content
        if text == company or text == 'QernelAI':
            # Company name - 72pt
            for run in para.runs:
                run.font.size = Pt(72)
                run.font.bold = True
                run.font.name = 'Calibri'
                # Set color based on text content
                if 'AI' in run.text:
                    run.font.color.rgb = hex_to_rgb(company_color)
                else:
                    run.font.color.rgb = hex_to_rgb('#000000')
        elif 'Engineering Specification' in text or 'Q32' in text:
            # Title line - 36pt
            for run in para.runs:
                run.font.size = Pt(36)
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.color.rgb = hex_to_rgb('#000000')
        elif text.startswith('Dec') or text.startswith('Jan') or text.startswith('20'):
            # Date - 36pt
            for run in para.runs:
                run.font.size = Pt(36)
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.color.rgb = hex_to_rgb('#000000')

    doc.save(str(docx_path))
    if not quiet:
        log("  Re-applied title page fonts")
    return True


def make_title_page(auto_token: str, primary_input: pathlib.Path) -> str:
    if auto_token != "__AUTO__":
        tp = pathlib.Path(auto_token)
        return read_text(tp)
    title = primary_input.stem.replace("_", " ").replace("-", " ").title()
    today = date.today().isoformat()
    return f"# {title}\n\n**Generated:** {today}\n\n"

def extract_latex_from_title_page(title_content: str) -> tuple[str, str]:
    """Extract raw LaTeX from title page content.

    Returns:
        tuple: (latex_content, remaining_markdown)
        If the title page is raw LaTeX, returns (latex, "")
        Otherwise returns ("", original_content)
    """
    # Check if the content is a raw LaTeX block
    latex_match = re.match(r'```\{=latex\}\s*\n(.*?)\n```\s*$', title_content, re.DOTALL)
    if latex_match:
        return latex_match.group(1), ""
    return "", title_content

def make_front_matter_lists(args) -> str:
    """Generate raw LaTeX commands for List of Tables/Figures/Waveforms.

    These are placed after the title page and before the main content.
    The TOC is auto-generated by Pandoc with --toc flag.
    Each list starts on a new page.
    """
    lists = []
    if args.lot:
        lists.append(r"\clearpage")
        lists.append(r"\listoftables")
    if args.lof:
        lists.append(r"\clearpage")
        lists.append(r"\listoffigures")
    if args.low:
        lists.append(r"\clearpage")
        lists.append(r"\listofwaveforms")
    if not lists:
        return ""
    # Wrap in raw LaTeX block for Pandoc
    return "\n```{=latex}\n" + "\n".join(lists) + "\n```\n\n"

def main():
    args = parse_args()
    in_path = pathlib.Path(args.input).resolve()
    out_docx = pathlib.Path(args.output).with_suffix(".docx")
    out_pdf  = out_docx.with_suffix(".pdf")

    with tempfile.TemporaryDirectory() as td:
        tmp_dir = pathlib.Path(td)
        build_md = tmp_dir / "build.md"
        tmp_imgs = tmp_dir / "gen_images"
        tmp_imgs.mkdir(exist_ok=True)

        # Determine source set
        if args.expand_index:
            files = collect_from_index(in_path, skip_index=args.skip_index_content)
            if not files:
                files = [in_path]
        else:
            files = [in_path]

        # Build merged markdown
        chunks = []
        title_page_latex = None  # For PDF: title page as raw LaTeX

        if args.title_page:
            title_content = make_title_page(args.title_page, in_path)
            # Extract LaTeX for PDF (title page before TOC)
            title_page_latex, remaining_md = extract_latex_from_title_page(title_content)
            if remaining_md:
                chunks.append(remaining_md)

        # Add List of Tables/Figures/Waveforms after title (appears after TOC in PDF)
        front_lists = make_front_matter_lists(args)
        if front_lists:
            chunks.append(front_lists)

        merged = concat_markdown(files, args.pagebreak)
        merged = strip_or_map_emoji(merged)
        merged = rewrite_wavedrom_images(merged, in_path.parent, tmp_imgs)
        if not args.no_wavedrom:
            merged = rewrite_wavedrom_blocks(merged, tmp_imgs, args.quiet, args.low)
        if not args.no_mermaid:
            merged = rewrite_mermaid_blocks(merged, tmp_imgs, args.quiet)

        chunks.append(merged)
        final_md = ("\n".join(chunks)).strip() + "\n"
        write_text(build_md, final_md)

        # Debug: save the merged markdown for inspection
        if args.debug_md:
            debug_md_path = pathlib.Path(args.output).with_suffix(".build.md")
            write_text(debug_md_path, final_md)
            log(f"Debug: saved merged markdown to {debug_md_path}")

        # Generate DOCX
        run_pandoc_docx(build_md, out_docx, args)

        # Apply corporate styling if --style specified
        if args.style:
            style_path = pathlib.Path(args.style)
            if style_path.exists():
                apply_docx_style(out_docx, style_path, args.quiet)
            else:
                log(f"  Warning: Style config not found: {args.style}")

        # Optional PDF
        if args.pdf:
            if args.style:
                # When --style is specified, generate PDF from styled DOCX using LibreOffice
                # This ensures the PDF matches the DOCX styling exactly
                soffice = shutil.which("soffice") or shutil.which("libreoffice")
                if not soffice:
                    log("  Warning: LibreOffice not found, falling back to Pandoc PDF")
                    run_pandoc_pdf(build_md, out_pdf, args, tmp_dir, title_page_latex)
                else:
                    # Update TOC/indexes before converting to PDF
                    if not args.quiet:
                        log(f"  Updating Table of Contents...")
                    update_docx_toc(out_docx, soffice, args.quiet)

                    # Re-apply title page styling after TOC update (LibreOffice may reset fonts)
                    if style_path.exists():
                        reapply_title_page_fonts(out_docx, style_path, args.quiet)

                    if not args.quiet:
                        log(f"  Converting styled DOCX to PDF using LibreOffice...")
                    subprocess.run([
                        soffice, "--headless", "--convert-to", "pdf",
                        "--outdir", str(out_pdf.parent), str(out_docx)
                    ], check=True)
            else:
                # Standard Pandoc PDF generation (LaTeX-based)
                try:
                    run_pandoc_pdf(build_md, out_pdf, args, tmp_dir, title_page_latex)
                except subprocess.CalledProcessError:
                    # Fallback: LibreOffice convert from DOCX, if present
                    soffice = shutil.which("soffice") or shutil.which("libreoffice")
                    if not soffice:
                        raise
                    subprocess.run([
                        soffice, "--headless", "--convert-to", "pdf",
                        "--outdir", str(out_pdf.parent), str(out_docx)
                    ], check=True)

    if not args.quiet:
        log(f"✓ Wrote {out_docx}")
        if args.pdf:
            log(f"✓ Wrote {out_pdf}")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)
